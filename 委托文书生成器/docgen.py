# -*- coding: utf-8 -*-
"""docx 模板填充模块：将【占位符】替换为字段值，尽量保留模板原格式。

遍历范围：正文段落、表格单元格、页眉、页脚。
占位符写法：【任意文字】（全角方括号），由"模板管理"页建立 占位符 → 数据字段 的映射。
"""
import copy
import re

from docx import Document
from docx.oxml.ns import qn

TOKEN_RE = re.compile(r"【[^】]{1,30}】")

# 证件号码组合标签：如"公民身份号码/统一社会信用代码"、"身份证号/统一社会信用代码"
# 生成时按当事人类型替换为单一标签（自然人→身份证号码；法人→统一社会信用代码）
_LABEL_COMBO_RE = re.compile(r"(?:公民身份号码|身份证号码|身份证号|身份号码)\s*/\s*统一社会信用代码")
# 组合标签 + 冒号且后面没有占位符（模板漏写【身份证号】的情形）→ 标签后直接补证件号
_LABEL_COMBO_TAIL_RE = re.compile(
    r"((?:公民身份号码|身份证号码|身份证号|身份号码)\s*/\s*统一社会信用代码)(\s*[:：][\s\u3000]*)$")


def _iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def extract_tokens(path):
    """提取模板中出现的全部【占位符】。"""
    doc = Document(path)
    tokens = set()
    for p in _iter_paragraphs(doc):
        full = "".join(r.text for r in p.runs)
        tokens.update(TOKEN_RE.findall(full))
    return sorted(tokens)


def _get_eastasia(run):
    """读取 run 的中文字体名（w:rFonts/@w:eastAsia），无则返回 None。"""
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return None
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        return None
    return rFonts.get(qn("w:eastAsia"))


def _append_run(p, text, rpr):
    """向段落 p 追加一个 run：格式复制自 rpr（可为 None），\t 保留为制表符。"""
    r = p._p.makeelement(qn("w:r"), {})
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    parts = text.split("\t")
    for i, part in enumerate(parts):
        if i > 0:
            r.append(r.makeelement(qn("w:tab"), {}))
        if part:
            t = r.makeelement(qn("w:t"), {})
            t.text = part
            t.set(qn("xml:space"), "preserve")
            r.append(t)
    p._p.append(r)


def _build_spans(full, token_map, values):
    """找出段落文本 full 中所有需要替换的区间，返回 [(start, end, 替换文本)]。

    覆盖三类：① 组合标签尾部（"身份证号/统一社会信用代码：" 段尾，补证号）
    ② 组合标签（"身份证号/统一社会信用代码" → 单一标签）
    ③ 【占位符】。区间按原文偏移给出，互不重叠。
    """
    spans = []
    label = values.get("client_id_label")
    id_no = values.get("client_id")

    if label:
        # ① 组合标签 + 冒号位于段尾 → 单一标签 + 冒号 + 证号
        if id_no:
            for m in _LABEL_COMBO_TAIL_RE.finditer(full):
                spans.append((m.start(), m.end(), str(label) + m.group(2) + str(id_no)))
        # ② 普通组合标签 → 单一标签（已被①覆盖的跳过）
        for m in _LABEL_COMBO_RE.finditer(full):
            covered = any(s <= m.start() and m.end() <= e for s, e, _ in spans)
            if not covered:
                spans.append((m.start(), m.end(), str(label)))
        # ③ 显式【证件号码类型】占位符
        for m in re.finditer(r"【证件号码类型】", full):
            spans.append((m.start(), m.end(), str(label)))

    # 普通【占位符】→ 字段值
    for m in TOKEN_RE.finditer(full):
        tok = m.group(0)
        key = token_map.get(tok)
        if not key:
            continue  # 未映射的占位符原样保留
        val = values.get(key)
        if val is None or val == "":
            continue  # 字段为空也保留占位符，便于人工补填
        spans.append((m.start(), m.end(), str(val)))
    return spans


def fill_template(template_path, token_map, values, out_path):
    """token_map: {占位符: 字段key}；values: {字段key: 字符串值}。

    替换策略：run 级替换 —— 只重建"占位符/组合标签所在"的 run，
    其余 run 原样保留，从而保留每个 run 的字体、加粗、制表符等局部格式。
    """
    doc = Document(template_path)

    replaced = 0
    for p in _iter_paragraphs(doc):
        runs = p.runs
        if not runs:
            continue
        full = "".join(r.text for r in runs)
        spans = _build_spans(full, token_map, values)
        if not spans:
            continue

        # 每个 run 在 full 中的字符区间
        run_spans = []
        pos = 0
        for r in runs:
            n = len(r.text)
            run_spans.append((pos, pos + n))
            pos += n

        # 受影响的 run 索引（与任一替换区间重叠）
        affected = set()
        for s, e, _ in spans:
            for ri, (rs, re_) in enumerate(run_spans):
                if rs < e and s < re_:
                    affected.add(ri)
        if not affected:
            continue

        # 重建整段：按"原文→替换"分段，未受影响的 run 原样保留
        # 先收集每个 run 的 rPr 格式模板
        rprs = []
        for r in runs:
            rPr = r._element.find(qn("w:rPr"))
            rprs.append(copy.deepcopy(rPr) if rPr is not None else None)

        spans_sorted = sorted(spans, key=lambda x: x[0])
        # 分段：光标在 full 上推进，span 之外的为"保留"段，span 内为"替换"段
        segments = []  # (text, 源run索引)
        cursor = 0
        for s, e, rep in spans_sorted:
            if s < cursor:
                continue  # 跳过重叠/已覆盖的 span
            if cursor < s:
                # 保留段 [cursor, s)，拆回原 run
                _append_equal_segments(segments, full, cursor, s, run_spans)
            # 替换段 [s, e)，用覆盖 s 的 run 的格式
            ri = next((i for i, (rs, re_) in enumerate(run_spans) if rs <= s < re_), 0)
            if rep != full[s:e]:
                segments.append((rep, ri))
            cursor = e
        if cursor < len(full):
            _append_equal_segments(segments, full, cursor, len(full), run_spans)

        # 合并相邻同源 run 的片段
        merged = []
        for txt, ri in segments:
            if not txt:
                continue
            if merged and merged[-1][1] == ri:
                merged[-1][0] += txt
            else:
                merged.append([txt, ri])

        # 删除所有原 run，按 merged 重建
        for r in runs:
            r._element.getparent().remove(r._element)
        for txt, ri in merged:
            _append_run(p, txt, rprs[ri])
        replaced += 1

    doc.save(out_path)
    return replaced


def _append_equal_segments(segments, full, start, end, run_spans):
    """把 full 的 [start, end) 保留段按原 run 边界切分，追加为 (原文, 源run索引) 片段。"""
    s = start
    for ri, (rs, re_) in enumerate(run_spans):
        if re_ <= s:
            continue
        if rs >= end:
            break
        cut = min(re_, end)
        if cut > s:
            segments.append((full[s:cut], ri))
        s = cut
        if s >= end:
            break


# ---------------- 诉讼地位角色（我方角色 ↔ 对方角色） ----------------
# 委托人不一定是原告：一审可能是被告，二审可能是被上诉人，再审可能是被申请人，
# 执行可能是被执行人。选定我方角色后，对方角色自动取反。
ROLE_PAIRS = {
    "原告": "被告", "被告": "原告",
    "上诉人": "被上诉人", "被上诉人": "上诉人",
    "再审申请人": "再审被申请人", "再审被申请人": "再审申请人",
    "申请执行人": "被执行人", "被执行人": "申请执行人",
    "申请人": "被申请人", "被申请人": "申请人",
}

# 程序阶段 → 该阶段可选的诉讼地位（前端按代理阶段联动过滤下拉）
STAGE_ROLES = {
    "一审": ["原告", "被告"],
    "二审": ["上诉人", "被上诉人"],
    "再审": ["再审申请人", "再审被申请人"],
    "再审申请": ["再审申请人", "再审被申请人"],
    "执行": ["申请执行人", "被执行人"],
    "仲裁": ["申请人", "被申请人"],
}


# ---------------- 字段定义（网页表单与模板映射共用） ----------------

FIELD_DEFS = [
    {"key": "client_name", "label": "委托人姓名/名称", "group": "both", "required": True},
    {"key": "client_gender", "label": "性别", "group": "person", "required": False},
    {"key": "client_ethnicity", "label": "民族", "group": "person", "required": False},
    {"key": "client_birth", "label": "出生日期", "group": "person", "required": False},
    {"key": "client_address", "label": "住址/住所", "group": "both", "required": True},
    {"key": "client_id", "label": "身份证号/统一社会信用代码", "group": "both", "required": True},
    {"key": "client_phone", "label": "委托人联系电话", "group": "both", "required": True},
    {"key": "opponent_name", "label": "对方当事人姓名/名称", "group": "both", "required": True},
    {"key": "client_role", "label": "我方当事人诉讼地位", "group": "both", "required": True},
    {"key": "case_reason", "label": "案由", "group": "both", "required": True},
    {"key": "agency_stage", "label": "代理阶段", "group": "both", "required": True},
    {"key": "auth_scope", "label": "代理权限", "group": "both", "required": True},
    {"key": "legal_rep", "label": "法定代表人姓名", "group": "company", "required": True},
    {"key": "legal_rep_duty", "label": "法定代表人职务", "group": "company", "required": True},
    {"key": "contract_no", "label": "合同编号", "group": "both", "required": False},
    {"key": "sign_date", "label": "签订日期", "group": "both", "required": False},
    {"key": "sign_year", "label": "日期·年", "group": "both", "required": False},
    {"key": "sign_month", "label": "日期·月", "group": "both", "required": False},
    {"key": "sign_day", "label": "日期·日", "group": "both", "required": False},
]

# 关键词 → 字段key 的自动映射规则（上传模板时预填建议）
# 注意顺序：【委托人角色/对方角色】须排在最前（"对方角色"含"对方"，不能先被"对方→被告"规则吃掉）；
# 【被告姓名】须先于【...姓名】；【法定代表人职务】须先于【法定代表人】；【法定代表人姓名】须先于【...姓名】
_AUTO_RULES = [
    (("委托人角色", "我方角色", "我方当事人角色", "诉讼地位"), "client_role"),
    (("对方角色", "对方当事人角色"), "opponent_role"),
    (("被告", "对方当事人", "对方"), "opponent_name"),
    (("职务",), "legal_rep_duty"),
    (("法定代表人", "法定代表"), "legal_rep"),
    (("姓名", "委托人名", "甲方名", "委托人名称"), "client_name"),
    (("性别",), "client_gender"),
    (("民族",), "client_ethnicity"),
    (("出生",), "client_birth"),
    (("住址", "住所", "地址", "营业场所"), "client_address"),
    (("身份证", "身份号码", "证号", "信用代码", "信用证号"), "client_id"),
    (("电话", "联系方式", "手机"), "client_phone"),
    (("案由", "纠纷", "案件"), "case_reason"),
    (("代理阶段", "审理阶段", "诉讼阶段"), "agency_stage"),
    (("权限", "授权"), "auth_scope"),
    (("编号", "合同号"), "contract_no"),
    (("签订日期", "签署日期", "日期"), "sign_date"),
]


def guess_field(token):
    """根据占位符文字猜测对应字段 key，猜不出返回 None。"""
    inner = token.strip("【】")
    if inner == "年":
        return "sign_year"
    if inner == "月":
        return "sign_month"
    if inner == "日":
        return "sign_day"
    if "证件号码类型" in inner:
        return None  # 由 fill_template 按当事人类型动态填充，无需映射
    for keywords, key in _AUTO_RULES:
        for kw in keywords:
            if kw in inner:
                return key
    return None


def _set_eastasia(run, font):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def fix_eastasia_fonts(path, body_font="仿宋", title_font="宋体"):
    """修复 textutil 转换 .doc→.docx 后中文字体标记丢失的问题。

    仅当 run 缺少中文字体（w:eastAsia 为空）时才补默认字体，绝不覆盖已有字体、
    字号、加粗等格式 —— 保证 WPS 另存的 .docx 等格式完整的模板不被改动。
    """
    doc = Document(path)
    title_done = False
    for p in _iter_paragraphs(doc):
        text = "".join(r.text for r in p.runs)
        if not text.strip():
            continue
        for r in p.runs:
            if _get_eastasia(r) is None:
                _set_eastasia(r, title_font if not title_done else body_font)
        if not title_done:
            title_done = True
    doc.save(path)


def inject_header_logo(path, logo_png, width_inches=2.4):
    """把律所 logo 注入 .docx 页眉首段（左对齐），用于 textutil 丢失 logo 的兜底。

    若页眉已含图片则跳过；logo 宽度默认 2.4 英寸，可据实际调整。
    """
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(path)
    section = doc.sections[0]
    header = section.header
    # 若页眉已有图片，则不重复注入
    if header.paragraphs:
        existing = "".join(p.text for p in header.paragraphs)
        if header.paragraphs[0]._element.findall(".//" + qn("w:drawing")):
            return False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(logo_png, width=Inches(width_inches))
    doc.save(path)
    return True
