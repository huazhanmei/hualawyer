# -*- coding: utf-8 -*-
"""首次运行时生成三份示例模板（含【...】占位符）。

格式遵循律所规范：标题宋体小二加粗、正文仿宋小三、一倍行距、律所页眉。
正式使用时请在"模板管理"页上传律所真实模板替换。
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

FIRM = "福建知信衡律师事务所"

TPL_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates_docx")


def _set_font(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _add_para(doc, text, size=15, bold=False, align=None, font="仿宋", indent=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    run = p.add_run(text)
    _set_font(run, font, size, bold)
    return p


def _add_header(doc):
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FIRM)
    _set_font(run, "宋体", 9)


def _base_doc():
    doc = Document()
    _add_header(doc)
    return doc


def _title(doc, text):
    _add_para(doc, text, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
              font="宋体", indent=False)
    _add_para(doc, "", size=15, font="仿宋", indent=False)


def make_contract():
    doc = _base_doc()
    _title(doc, "委托代理合同")
    _add_para(doc, "合同编号：【合同编号】", indent=False)
    _add_para(doc, "甲方（委托人）：【委托人姓名】")
    _add_para(doc, "【证件号码类型】：【身份证号】")
    _add_para(doc, "住址/住所：【住址】")
    _add_para(doc, "联系电话：【联系电话】")
    _add_para(doc, "乙方（受托人）：%s" % FIRM)
    _add_para(doc, "鉴于甲方因与【被告姓名】【案由】纠纷一案需要委托律师提供法律服务，"
                   "根据《中华人民共和国民法典》《中华人民共和国律师法》等有关规定，"
                   "双方经协商一致，订立本合同。")
    _add_para(doc, "第一条　委托事项")
    _add_para(doc, "甲方因与【被告姓名】【案由】纠纷一案，"
                   "委托乙方指派律师在【代理阶段】阶段担任甲方的委托代理人。")
    _add_para(doc, "第二条　代理权限")
    _add_para(doc, "乙方律师的代理权限为：【代理权限】。")
    _add_para(doc, "第三条　代理费及支付方式")
    _add_para(doc, "双方另行协商确定，以收费协议为准。")
    _add_para(doc, "第四条　合同生效")
    _add_para(doc, "本合同自双方签字（盖章）之日起生效。")
    _add_para(doc, "")
    _add_para(doc, "甲方：【委托人姓名】　　　　乙方：%s" % FIRM, indent=False)
    _add_para(doc, "【年】年【月】月【日】日　　　　【年】年【月】月【日】日", indent=False)
    path = os.path.join(TPL_DIR, "示例-委托代理合同.docx")
    doc.save(path)
    return path


def make_authorization():
    doc = _base_doc()
    _title(doc, "授权委托书")
    _add_para(doc, "委托人：【委托人姓名】")
    _add_para(doc, "【证件号码类型】：【身份证号】")
    _add_para(doc, "住址/住所：【住址】")
    _add_para(doc, "受托人：%s　律师" % FIRM)
    _add_para(doc, "原告【委托人姓名】与被告【被告姓名】【案由】纠纷一案，"
                   "现委托上述受托人在【代理阶段】阶段作为我方委托诉讼代理人。")
    _add_para(doc, "代理权限：【代理权限】。")
    _add_para(doc, "")
    _add_para(doc, "委托人：【委托人姓名】", indent=False)
    _add_para(doc, "【年】年【月】月【日】日", indent=False)
    path = os.path.join(TPL_DIR, "示例-授权委托书.docx")
    doc.save(path)
    return path


def make_legal_rep():
    doc = _base_doc()
    _title(doc, "法定代表人身份证明")
    _add_para(doc, "兹证明【法定代表人姓名】在我司任【法定代表人职务】职务，"
                   "系我司（【委托人名称】，统一社会信用代码：【统一社会信用代码】）的法定代表人。")
    _add_para(doc, "单位住所：【委托人住所】")
    _add_para(doc, "特此证明。")
    _add_para(doc, "")
    _add_para(doc, "【年】年【月】月【日】日", indent=False)
    _add_para(doc, "公司（盖章）：【委托人名称】", align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    path = os.path.join(TPL_DIR, "示例-法定代表人身份证明.docx")
    doc.save(path)
    return path


def ensure_demo_templates():
    os.makedirs(TPL_DIR, exist_ok=True)
    made = []
    if not any(f.endswith(".docx") for f in os.listdir(TPL_DIR)):
        made.append(make_contract())
        made.append(make_authorization())
        made.append(make_legal_rep())
    return made
